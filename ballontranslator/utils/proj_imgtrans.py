import hashlib
import io
import os, json, shutil, re, docx, docx2txt, piexif, cv2
import tempfile
import warnings
from docx.shared import Inches
from docx import Document
import piexif.helper
import numpy as np
import os.path as osp
from typing import BinaryIO, Optional, Tuple, Union, List, Dict
from PIL import Image

from .logger import logger as LOGGER
from .io_utils import find_all_imgs, imread, imwrite, NumpyEncoder
from .textblock import (
    FontFormat,
    TextBlock,
    normalize_textblock_effect_payload,
)
from .fontformat import warn_ignored_legacy_effects
from .raster_assets import RasterAssetRef
from .rgba import premultiply_rgba_in_place
from .text_alpha_mask import AlphaBrushStroke, TextAlphaMask
from .config import pcfg, RunStatus
from . import shared

class ProjectDirNotExistException(Exception):
    pass

class ProjectLoadFailureException(Exception):
    pass

class ProjectNotSupportedException(Exception):
    pass

class ImgnameNotInProjectException(Exception):
    pass


RASTER_ASSET_MAX_SOURCE_BYTES = 32 * 1024 * 1024
RASTER_ASSET_MAX_PIXELS = 64 * 1024 * 1024
RASTER_ASSET_MAX_DECODED_BYTES = RASTER_ASSET_MAX_PIXELS * 4
RASTER_ASSET_DECODE_CACHE_ITEMS = 2
RASTER_ASSET_DECODE_CACHE_MAX_BYTES = RASTER_ASSET_MAX_DECODED_BYTES * 2
_RASTER_ASSET_RGBA8_MODES = {
    '1', 'L', 'LA', 'P', 'RGB', 'RGBA', 'CMYK', 'YCbCr', 'HSV'
}


def get_last_modified_file(file_prefix, exts, ext_fallback=None):
    '''
    get last modified file from files sharing same prefix
    '''
    latest_time = -1
    latest_f = None
    for ext in exts:
        tmp_p = file_prefix + ext
        if osp.exists(tmp_p) and osp.getmtime(tmp_p) > latest_time:
            latest_time = osp.getmtime(tmp_p)
            latest_f = tmp_p
    if latest_f is None:
        if ext_fallback is not None:
            latest_f = file_prefix + ext_fallback
        else:
            latest_f = file_prefix + exts[0]
    return latest_f


def write_jpg_metadata(imgpath: str, metadata="a metadata"):
    exif_dict = {"Exif":{piexif.ExifIFD.UserComment: piexif.helper.UserComment.dump(metadata, encoding='unicode')}}
    exif_bytes = piexif.dump(exif_dict)
    piexif.insert(exif_bytes, imgpath)

def read_jpg_metadata(imgpath: str):
    exif_dict = piexif.load(imgpath)
    user_comment = piexif.helper.UserComment.load(exif_dict["Exif"][piexif.ExifIFD.UserComment])
    bubdict = json.loads(user_comment)
    return bubdict

page_start_pattern = re.compile(r'^###\s+', re.MULTILINE)
text_blkid_start_pattern = re.compile(r'^\d+\.', re.MULTILINE)

def parse_txt_translation(file_path: str):
    with open(file_path, 'r', encoding='utf8') as f:
        content = f.read()
    page_start = None
    page_list = []
    for matched in page_start_pattern.finditer(content):
        start, end = matched.span()
        if page_start is not None:
            page_list.append({'page_content': content[page_start: start]})
        page_start = start
    if page_start is not None:
        page_list.append({'page_content': content[page_start:]})

    for page_dict in page_list:
        page_content = page_dict['page_content']
        page_dict['page_name'] = page_start_pattern.sub('', page_content.split('\n')[0]).strip()
        blkid_start = blkid_end = None
        blk_list = []
        for matched in text_blkid_start_pattern.finditer(page_content):
            start, end = matched.span()
            if blkid_start is not None:
                blk_list.append(page_content[blkid_end: start].strip())
            blkid_start = start
            blkid_end = end
        if blkid_start is not None:
            blk_list.append(page_content[blkid_end:].strip())
        page_dict['blk_list'] = blk_list

    return page_list


class TextBlkEncoder(NumpyEncoder):
    def default(self, obj):
        if isinstance(obj, TextBlock):
            return obj.to_dict()
        elif isinstance(obj, FontFormat):
            return obj.to_serializable_dict()
        elif isinstance(obj, (TextAlphaMask, AlphaBrushStroke)):
            return obj.to_serializable_dict()
        return NumpyEncoder.default(self, obj)


class ProjImgTrans:

    def __init__(self, directory: str = None):
        self._load_identity = object()
        self._raster_asset_cache: Dict[
            str,
            Tuple[
                Tuple[int, int, int, int],
                np.ndarray,
                Optional[np.ndarray],
            ],
        ] = {}
        self.type = 'imgtrans'
        self.directory: str = None
        self.pages: Dict[str, List[TextBlock]] = {}
        self._pagename2idx = {}
        self._idx2pagename = {}
        self._image_info = {}

        self._fuzzy_inpainted_list = None

        self.not_found_pages: Dict[str, List[TextBlock]] = {}
        self.new_pages: List[str] = []
        self.proj_path: str = None

        self.current_img: str = None
        self.img_array: np.ndarray = None
        self.mask_array: np.ndarray = None
        self.inpainted_array: np.ndarray = None
        if directory is not None:
            self.load(directory)

    @property
    def load_identity(self):
        """Return the opaque identity of the currently loaded project contents.

        >>> project = ProjImgTrans()
        >>> project.load_identity is project.load_identity
        True
        """
        return self._load_identity

    def idx2pagename(self, idx: int) -> str:
        return self._idx2pagename[idx]

    def pagename2idx(self, pagename: str) -> int:
        if pagename in self.pages:
            return self._pagename2idx[pagename]
        return -1

    def proj_name(self) -> str:
        return self.type+'_'+osp.basename(self.directory)

    def load(self, directory: str, json_path: str = None) -> bool:
        self._raster_asset_cache.clear()
        self.directory = directory
        if json_path is None:
            self.proj_path = osp.join(
                self.directory, self.proj_name() + '.json'
            )
        else:
            self.proj_path = json_path
        new_proj = False
        if not osp.exists(self.proj_path):
            new_proj = True
            self.new_project()
        else:
            try:
                with open(self.proj_path, 'r', encoding='utf8') as f:
                    proj_dict = json.loads(f.read())
            except Exception as e:
                raise ProjectLoadFailureException(e)
            self.load_from_dict(proj_dict)
        if not osp.exists(self.inpainted_dir()):
            os.makedirs(self.inpainted_dir())
        if not osp.exists(self.mask_dir()):
            os.makedirs(self.mask_dir())

        return new_proj

    def mask_dir(self):
        return osp.join(self.directory, 'mask')

    def inpainted_dir(self):
        return osp.join(self.directory, 'inpainted')

    def result_dir(self):
        return osp.join(self.directory, 'result')

    def assets_dir(self) -> str:
        """Return the project-owned directory for reusable immutable assets."""
        return osp.join(self.directory, 'assets')

    def _resolved_assets_root(self, *, create: bool = False) -> str:
        """Return the real assets root only when it stays in the project."""
        if not self.directory or not osp.isdir(self.directory):
            raise ProjectDirNotExistException
        assets_dir = self.assets_dir()
        if create:
            os.makedirs(assets_dir, exist_ok=True)
        project_root = osp.realpath(self.directory)
        assets_root = osp.realpath(assets_dir)
        try:
            contained = (
                osp.commonpath((project_root, assets_root)) == project_root
            )
        except ValueError:
            contained = False
        if not contained:
            raise OSError('project assets directory resolves outside the project')
        return assets_root

    @staticmethod
    def _raster_extension(image_format: str) -> str:
        extensions = {
            'BMP': '.bmp',
            'JPEG': '.jpg',
            'JPEGXL': '.jxl',
            'JXL': '.jxl',
            'PNG': '.png',
            'WEBP': '.webp',
        }
        try:
            return extensions[image_format.upper()]
        except (AttributeError, KeyError) as error:
            raise ValueError('unsupported raster asset format') from error

    @staticmethod
    def _hash_raster_asset_file(path: str) -> str:
        """Hash one size-bounded raster file.

        >>> len(ProjImgTrans._hash_raster_asset_file(__file__))
        64
        """
        digest = hashlib.sha256()
        total = 0
        with open(path, 'rb') as source:
            for chunk in iter(lambda: source.read(1024 * 1024), b''):
                total += len(chunk)
                if total > RASTER_ASSET_MAX_SOURCE_BYTES:
                    raise ValueError('raster asset exceeds the source-byte limit')
                digest.update(chunk)
        return digest.hexdigest()

    @classmethod
    def _decode_raster_asset_snapshot(
        cls, path: str
    ) -> Tuple[str, np.ndarray]:
        """Fully decode one bounded snapshot to owned immutable RGBA8 pixels.

        Source bytes are capped at 32 MiB, images at 64 Mpx, and decoded RGBA
        storage at 256 MiB. Integer/float images wider than eight bits are
        rejected rather than truncated.

        >>> callable(ProjImgTrans._decode_raster_asset_snapshot)
        True
        """
        if osp.getsize(path) > RASTER_ASSET_MAX_SOURCE_BYTES:
            raise ValueError('raster asset exceeds the source-byte limit')
        try:
            with warnings.catch_warnings():
                warnings.simplefilter('error', Image.DecompressionBombWarning)
                with Image.open(path) as image:
                    extension = cls._raster_extension(image.format)
                    width, height = image.size
                    pixels = width * height
                    if pixels <= 0 or pixels > RASTER_ASSET_MAX_PIXELS:
                        raise ValueError('raster asset exceeds the pixel limit')
                    if pixels * 4 > RASTER_ASSET_MAX_DECODED_BYTES:
                        raise ValueError(
                            'raster asset exceeds the decoded-byte limit'
                        )
                    if image.mode not in _RASTER_ASSET_RGBA8_MODES:
                        raise ValueError(
                            'raster asset must use supported 8-bit channels'
                        )
                    rgba = np.array(image.convert('RGBA'), dtype=np.uint8)
        except ValueError:
            raise
        except (
            Image.DecompressionBombError,
            Image.DecompressionBombWarning,
            MemoryError,
            OSError,
        ) as error:
            raise ValueError('unable to decode raster asset') from error
        if rgba.shape != (height, width, 4):
            raise ValueError('unable to decode raster asset as RGBA8')
        rgba = np.ascontiguousarray(rgba)
        rgba.flags.writeable = False
        return extension, rgba

    @staticmethod
    def _raster_asset_signature(path: str) -> Tuple[int, int, int, int]:
        stat = os.stat(path)
        return stat.st_ino, stat.st_size, stat.st_mtime_ns, stat.st_ctime_ns

    def _cache_raster_asset(
        self,
        asset: RasterAssetRef,
        rgba: np.ndarray,
        signature: Tuple[int, int, int, int],
    ) -> None:
        key = asset.path
        self._raster_asset_cache.pop(key, None)
        self._raster_asset_cache[key] = (signature, rgba, None)
        self._trim_raster_asset_cache(key)

    @staticmethod
    def _raster_asset_cache_entry_bytes(
        cached: Tuple[
            Tuple[int, int, int, int],
            np.ndarray,
            Optional[np.ndarray],
        ],
    ) -> int:
        """Count unique straight and premultiplied storage in one entry.

        >>> rgba = np.zeros((2, 3, 4), dtype=np.uint8)
        >>> cached = ((0, 0, 0, 0), rgba, rgba)
        >>> ProjImgTrans._raster_asset_cache_entry_bytes(cached)
        24
        """
        rgba, premultiplied = cached[1], cached[2]
        return rgba.nbytes + (
            premultiplied.nbytes
            if premultiplied is not None and premultiplied is not rgba
            else 0
        )

    def _trim_raster_asset_cache(self, retained_path: str) -> None:
        """Retain the requested asset while enforcing count and byte bounds.

        >>> callable(ProjImgTrans._trim_raster_asset_cache)
        True
        """
        while len(self._raster_asset_cache) > 1:
            cached_bytes = sum(
                self._raster_asset_cache_entry_bytes(cached)
                for cached in self._raster_asset_cache.values()
            )
            if (
                len(self._raster_asset_cache)
                <= RASTER_ASSET_DECODE_CACHE_ITEMS
                and cached_bytes <= RASTER_ASSET_DECODE_CACHE_MAX_BYTES
            ):
                break
            oldest = next(
                path for path in self._raster_asset_cache
                if path != retained_path
            )
            self._raster_asset_cache.pop(oldest)

    def _cached_raster_asset_pixels(
        self,
        asset: RasterAssetRef,
        cached: Tuple[
            Tuple[int, int, int, int],
            np.ndarray,
            Optional[np.ndarray],
        ],
        *,
        premultiplied: bool,
    ) -> np.ndarray:
        """Return the requested immutable representation from one cache entry.

        >>> callable(ProjImgTrans._cached_raster_asset_pixels)
        True
        """
        if not premultiplied:
            return cached[1]
        if cached[2] is None:
            rgba = cached[1]
            if np.all(rgba[..., 3] == 255):
                premultiplied_rgba = rgba
            else:
                premultiplied_rgba = np.array(rgba, copy=True, order='C')
                premultiply_rgba_in_place(premultiplied_rgba)
                premultiplied_rgba.flags.writeable = False
            cached = (cached[0], rgba, premultiplied_rgba)
            self._raster_asset_cache[asset.path] = cached
            self._trim_raster_asset_cache(asset.path)
        assert cached[2] is not None
        return cached[2]

    def _import_raster_asset_stream(
        self,
        source: BinaryIO,
        display_name: str,
    ) -> RasterAssetRef:
        """Validate and atomically install one already-open raster stream."""
        assets_dir = self._resolved_assets_root(create=True)
        temporary_path = None
        try:
            digest = hashlib.sha256()
            total = 0
            with tempfile.NamedTemporaryFile(
                mode='wb', prefix='.import-', dir=assets_dir, delete=False
            ) as temporary:
                temporary_path = temporary.name
                for chunk in iter(lambda: source.read(1024 * 1024), b''):
                    total += len(chunk)
                    if total > RASTER_ASSET_MAX_SOURCE_BYTES:
                        raise ValueError(
                            'raster asset exceeds the source-byte limit'
                        )
                    digest.update(chunk)
                    temporary.write(chunk)
            extension, rgba = self._decode_raster_asset_snapshot(
                temporary_path
            )
            digest_hex = digest.hexdigest()
            filename = digest_hex + extension
            destination = osp.join(assets_dir, filename)
            if (
                osp.commonpath((assets_dir, osp.realpath(destination)))
                != assets_dir
            ):
                raise OSError(
                    'raster asset destination resolves outside assets'
                )
            if osp.exists(destination):
                if self._hash_raster_asset_file(destination) != digest_hex:
                    raise OSError(
                        'content-addressed raster asset has unexpected contents'
                    )
            else:
                os.replace(temporary_path, destination)
                temporary_path = None
            asset = RasterAssetRef(
                f'assets/{filename}', osp.basename(display_name)
            )
            self._cache_raster_asset(
                asset, rgba, self._raster_asset_signature(destination)
            )
            return asset
        except (
            Image.DecompressionBombError,
            Image.DecompressionBombWarning,
            MemoryError,
        ) as error:
            raise ValueError('unable to read raster asset') from error
        finally:
            if temporary_path is not None and osp.exists(temporary_path):
                os.unlink(temporary_path)

    def import_raster_asset(self, source_path: str) -> RasterAssetRef:
        """Snapshot and validate one raster before content-addressing it.

        The source is opened once. Hashing and decoding both use the bounded
        temporary copy, so a concurrently replaced source cannot split the
        identity from the pixels that were validated.

        >>> callable(ProjImgTrans.import_raster_asset)
        True
        """
        if not isinstance(source_path, str) or not osp.isfile(source_path):
            raise ValueError('raster asset source must be an existing file')
        with open(source_path, 'rb') as source:
            return self._import_raster_asset_stream(source, source_path)

    def import_raster_asset_bytes(
        self,
        payload: bytes,
        display_name: str = 'generated.png',
    ) -> RasterAssetRef:
        """Import generated raster bytes through the normal asset boundary.

        >>> callable(ProjImgTrans.import_raster_asset_bytes)
        True
        """
        if not isinstance(payload, bytes) or not payload:
            raise ValueError('raster asset payload must be non-empty bytes')
        if not isinstance(display_name, str) or not display_name.strip():
            raise ValueError('raster asset display name must be non-empty')
        return self._import_raster_asset_stream(
            io.BytesIO(payload), display_name
        )

    def _resolve_raster_asset_path(
        self,
        asset: RasterAssetRef,
    ) -> Optional[str]:
        """Resolve containment and existence without reading asset bytes."""
        try:
            assets_root = self._resolved_assets_root()
        except (OSError, ProjectDirNotExistException):
            return None
        path = osp.realpath(osp.join(self.directory, *asset.path.split('/')))
        try:
            contained = osp.commonpath((assets_root, path)) == assets_root
        except ValueError:
            contained = False
        return path if contained and osp.isfile(path) else None

    def resolve_raster_asset(
        self,
        asset: RasterAssetRef,
        *,
        strict: bool = False,
    ) -> Optional[str]:
        """Resolve a managed relative reference without escaping ``assets``.

        Missing optional assets are bypassed during interactive rendering;
        strict export receives an exception instead.
        """
        if not isinstance(asset, RasterAssetRef):
            raise TypeError('raster asset resolution requires RasterAssetRef')
        path = self._resolve_raster_asset_path(asset)
        if path is not None:
            if strict:
                signature = self._raster_asset_signature(path)
                digest = self._hash_raster_asset_file(path)
                if self._raster_asset_signature(path) != signature:
                    raise OSError('raster asset changed while it was verified')
                if digest != asset.digest:
                    raise OSError(
                        f'Raster asset contents do not match: {asset.path}'
                    )
            return path
        message = f'Raster asset is unavailable: {asset.path}'
        if strict:
            raise FileNotFoundError(message)
        LOGGER.warning(message)
        return None

    def load_raster_asset(
        self,
        asset: RasterAssetRef,
        *,
        strict: bool = False,
        premultiplied: bool = False,
    ) -> Optional[np.ndarray]:
        """Return shared immutable straight or premultiplied RGBA8 pixels.

        Every read resolves existence and containment before a decoded-cache
        hit. Unchanged warm entries stay hash-free; cold or stat-changed files
        are digest-verified before decoding. Strict reads always verify source
        bytes inside the same signature bracket as cache reuse or decode.
        Failures are not cached, so restore plus invalidation recovers.
        """
        if not isinstance(asset, RasterAssetRef):
            raise TypeError('raster asset loading requires RasterAssetRef')
        path = self._resolve_raster_asset_path(asset)
        if path is None:
            message = f'Raster asset is unavailable: {asset.path}'
            if strict:
                raise FileNotFoundError(message)
            LOGGER.warning(message)
            return None
        try:
            signature = self._raster_asset_signature(path)
            cached = self._raster_asset_cache.get(asset.path)
            warm = cached is not None and cached[0] == signature
            if warm and not strict:
                return self._cached_raster_asset_pixels(
                    asset, cached, premultiplied=premultiplied
                )
            if not warm:
                self._raster_asset_cache.pop(asset.path, None)
            digest = self._hash_raster_asset_file(path)
            if digest != asset.digest:
                raise OSError(
                    f'Raster asset contents do not match: {asset.path}'
                )
            if warm:
                assert cached is not None
                rgba = self._cached_raster_asset_pixels(
                    asset, cached, premultiplied=premultiplied
                )
            else:
                extension, rgba = self._decode_raster_asset_snapshot(path)
                if not asset.path.endswith(extension):
                    raise ValueError(
                        'raster asset format does not match its ref'
                    )
            if self._raster_asset_signature(path) != signature:
                raise OSError('raster asset changed while it was loading')
            if warm:
                return rgba
        except (OSError, TypeError, ValueError) as error:
            self._raster_asset_cache.pop(asset.path, None)
            if strict:
                raise
            LOGGER.warning('Unable to decode Raster asset: %s', error)
            return None
        self._cache_raster_asset(asset, rgba, signature)
        cached = self._raster_asset_cache[asset.path]
        return self._cached_raster_asset_pixels(
            asset, cached, premultiplied=premultiplied
        )

    def load_from_dict(self, proj_dict: dict):
        self._raster_asset_cache.clear()
        self.set_current_img(None)
        effect_notices = set()

        def load_blocks(records: List[dict]) -> List[TextBlock]:
            blocks = []
            for record in records:
                normalized, notices = normalize_textblock_effect_payload(
                    record
                )
                effect_notices.update(notices)
                blocks.append(TextBlock(**normalized))
            return blocks

        try:
            self.pages = {}
            self._pagename2idx = {}
            self._idx2pagename = {}
            self.not_found_pages = {}
            page_dict = proj_dict['pages']
            not_found_pages = list(page_dict.keys())
            found_pages = find_all_imgs(img_dir=self.directory, abs_path=False, sort=True)
            for ii, imname in enumerate(found_pages):
                if imname in page_dict:
                    self.pages[imname] = load_blocks(page_dict[imname])
                    not_found_pages.remove(imname)
                else:
                    self.pages[imname] = []
                    self.new_pages.append(imname)
                self._pagename2idx[imname] = ii
                self._idx2pagename[ii] = imname
            for imname in not_found_pages:
                self.not_found_pages[imname] = load_blocks(page_dict[imname])
        except Exception as e:
            raise ProjectNotSupportedException(e)
        warn_ignored_legacy_effects(effect_notices, 'project')
        
        if 'image_info' in proj_dict:
            self._image_info = proj_dict['image_info']
        else:
            self._image_info = {}

        for p in self.pages:
            if p not in self._image_info:
                self._image_info[p] = {}
            img_info = self._image_info[p]
            if 'finish_code' not in img_info:
                page_blklist = self.pages[p]
                has_empty_blk = len(page_blklist) == 0 or \
                    any(not blk.text or len(blk.text) == 0 for blk in page_blklist)
                if has_empty_blk:
                    img_info['finish_code'] = 0
                else:
                    img_info['finish_code'] = RunStatus.FIN_ALL
            
        set_img_failed = False
        if 'current_img' in proj_dict:
            current_img = proj_dict['current_img']
            try:
                self.set_current_img(current_img)
            except ImgnameNotInProjectException:
                set_img_failed = True
        else:
            set_img_failed = True

        if set_img_failed:
            if len(self.pages) > 0:
                self.set_current_img_byidx(0)
        self._load_identity = object()

    def get_page_progress(self, pagename: str):
        fin_code = self._image_info[pagename]['finish_code']
        return (fin_code & pcfg.module.finish_code) == pcfg.module.finish_code

    def set_page_progress(self, pagename, code):
        self._image_info[pagename]['finish_code'] = code
        if not (code & RunStatus.FIN_TRANSLATE):
            self._image_info[pagename].pop('translation_target', None)

    def clear_page_progress(self, pagename, code):
        self._image_info[pagename]['finish_code'] &= ~code
        if code & RunStatus.FIN_TRANSLATE:
            self._image_info[pagename].pop('translation_target', None)

    def update_page_progress(self, pagename, code):
        self._image_info[pagename]['finish_code'] |= code

    def invalidate_translation(self, page_key):
        self.clear_page_progress(page_key, RunStatus.FIN_TRANSLATE)

    def begin_detection(self, page_key):
        """Invalidate translation before detection can replace page blocks."""
        self.invalidate_translation(page_key)

    def begin_full_page_translation(self, page_key):
        """Invalidate old completion until a full translation succeeds."""
        self.invalidate_translation(page_key)

    def mark_translation_finished(self, page_key, target_language):
        self.update_page_progress(page_key, RunStatus.FIN_TRANSLATE)
        self._image_info[page_key]['translation_target'] = target_language

    def load_translation_from_txt(self, file_path: str, target_language=None):
        page_list = parse_txt_translation(file_path)
        missing_pages = []
        unmatched_pages = []
        unexpected_pages = []
        matched_pages = []
        for page_dict in page_list:
            page_name = page_dict['page_name']
            if page_name in self.pages:
                matched_pages.append(page_name)
            else:
                unexpected_pages.append(page_name)
                continue
            blklist = self.pages[page_name]
            n_blk = len(blklist)
            src_blk_list = page_dict['blk_list']
            n_src_blk = len(src_blk_list)
            if n_src_blk != n_blk:
                LOGGER.warning(f'Unmatched text blocks in {page_name}, number of text blocks in this page vs source file: {n_blk}-{n_src_blk}')
                unmatched_pages.append(page_name)
            for blkid in range(min(n_blk, n_src_blk)):
                blk = blklist[blkid]
                blk.rich_text = ''
                blk.translation = src_blk_list[blkid]

        matched_pages = set(matched_pages)
        if len(matched_pages) != self.num_pages:
            for page_name in self.pages:
                if page_name not in matched_pages:
                    missing_pages.append(page_name)
        
        all_matched = (
            len(missing_pages) == 0
            and len(unmatched_pages) == 0
            and len(unexpected_pages) == 0
        )
        unmatched_page_set = set(unmatched_pages)
        for page_name in matched_pages - unmatched_page_set:
            if target_language is None:
                self.update_page_progress(page_name, RunStatus.FIN_TRANSLATE)
                self._image_info[page_name].pop('translation_target', None)
            else:
                self.mark_translation_finished(page_name, target_language)

        # Completion is page-specific: malformed imported pages are invalidated,
        # while project pages absent from the import retain their existing state.
        for page_name in unmatched_page_set:
            self.clear_page_progress(page_name, RunStatus.FIN_TRANSLATE)
        return all_matched, {
            'missing_pages': missing_pages,
            'unmatched_pages': unmatched_pages,
            'unexpected_pages': unexpected_pages,
            'matched_pages': matched_pages,
        }

    def load_from_json(self, json_path: str):
        old_dir = self.directory
        directory = osp.dirname(json_path)
        try:
            self.load(directory, json_path=json_path)
        except Exception as e:
            self.load(old_dir)
            raise ProjectLoadFailureException(e)

    def set_current_img(self, imgname: str):
        if imgname is not None:
            if imgname not in self.pages:
                raise ImgnameNotInProjectException
            self.current_img = imgname
            img_path = self.current_img_path()
            mask_path = self.get_mask_path(get_last_modified=True)
            self.img_array = imread(img_path)
            im_h, im_w = self.img_array.shape[:2]
            if osp.exists(mask_path):
                self.mask_array = imread(mask_path, cv2.IMREAD_GRAYSCALE)
            else:
                self.mask_array = np.zeros((im_h, im_w), dtype=np.uint8)
            self.inpainted_array = self.load_inpainted_by_imgname(imgname)
            if self.inpainted_array is None:
                self.inpainted_array = np.copy(self.img_array)
        else:
            self.current_img = None
            self.img_array = None
            self.mask_array = None
            self.inpainted_array = None

    def current_has_alpha(self):
        if self.current_img is None:
            return False
        return len(self.img_array.shape) and self.img_array.shape[-1] == 4

    def set_current_img_byidx(self, idx: int):
        num_pages = self.num_pages
        if idx < 0:
            idx = idx + self.num_pages
        if idx < 0 or idx > num_pages - 1:
            self.set_current_img(None)
        else:
            self.set_current_img(self.idx2pagename(idx))

    def get_blklist_byidx(self, idx: int) -> List[TextBlock]:
        return self.pages[self.idx2pagename(idx)]

    @property
    def num_pages(self) -> int:
        return len(self.pages)

    @property
    def current_idx(self) -> int:
        return self.pagename2idx(self.current_img)

    def new_project(self):
        if not osp.exists(self.directory):
            raise ProjectDirNotExistException
        self._raster_asset_cache.clear()
        self.set_current_img(None)
        imglist = find_all_imgs(self.directory, abs_path=False, sort=True)
        self.pages = {}
        self._pagename2idx = {}
        self._idx2pagename = {}
        self._image_info = {}
        for ii, imgname in enumerate(imglist):
            self.pages[imgname] = []
            self._pagename2idx[imgname] = ii
            self._idx2pagename[ii] = imgname
            self._image_info[imgname] = {'finish_code': 0}
        self.set_current_img_byidx(0)
        self.save()
        self._load_identity = object()
        
    def save(self, keep_exist_as_backup=False):
        if not osp.exists(self.directory):
            raise ProjectDirNotExistException
        tmp_save_tgt = self.proj_path + '.tmp'
        try:
            with open(tmp_save_tgt, "w", encoding="utf-8") as f:
                f.write(json.dumps(self.to_dict(), ensure_ascii=False, cls=TextBlkEncoder))
        except:
            raise Exception(f'Failed to write {self.to_dict()}')
        if osp.exists(self.proj_path) and keep_exist_as_backup:
            os.replace(self.proj_path, self.proj_path + '.backup')
            os.replace(tmp_save_tgt, self.proj_path)
        else:
            os.replace(tmp_save_tgt, self.proj_path)
        LOGGER.debug(f'project saved')

    def to_dict(self) -> Dict:
        pages = self.pages.copy()
        pages.update(self.not_found_pages)        
        image_info = self._image_info.copy()
        return {
            'directory': self.directory,
            'pages': pages,
            'current_img': self.current_img,
            'image_info': image_info,
        }

    def read_img(self, imgname: str) -> np.ndarray:
        if imgname not in self.pages:
            raise ImgnameNotInProjectException
        img_path = osp.join(self.directory, imgname)
        img = imread(img_path)
        h, w = img.shape[:2]
        self._image_info[imgname].update({'width': w, 'height': h})
        return img

    def save_mask(self, img_name, mask: np.ndarray):
        imwrite(self.get_mask_path(img_name), mask, ext=pcfg.intermediate_imgsave_ext)

    def save_inpainted(self, img_name, inpainted: np.ndarray):
        imwrite(self.get_inpainted_path(img_name), inpainted, ext=pcfg.intermediate_imgsave_ext)

    def current_img_path(self) -> str:
        if self.current_img is None:
            return None
        return osp.join(self.directory, self.current_img)

    def get_mask_path(self, imgname: str = None, get_last_modified=False) -> str:
        if imgname is None:
            imgname = self.current_img

        fileprefix = osp.join(self.mask_dir(), osp.splitext(imgname)[0])
        if get_last_modified:
            p = get_last_modified_file(fileprefix, ['.jxl', '.png'], ext_fallback=pcfg.intermediate_imgsave_ext)
        else:
            p = fileprefix+pcfg.intermediate_imgsave_ext

        return p
    
    def load_mask_by_imgname(self, imgname: str) -> np.ndarray:
        mask = None
        mp = self.get_mask_path(imgname, get_last_modified=True)
        if osp.exists(mp):
            mask = imread(mp, cv2.IMREAD_GRAYSCALE)
        return mask

    def get_inpainted_path(self, imgname: str = None, get_last_modified=False) -> str:
        if imgname is None:
            imgname = self.current_img

        fileprefix = osp.join(self.inpainted_dir(), osp.splitext(imgname)[0])
        if get_last_modified:
            p = get_last_modified_file(fileprefix, ['.jxl', '.png'], ext_fallback=pcfg.intermediate_imgsave_ext)
        else:
            p = fileprefix+pcfg.intermediate_imgsave_ext

        if not osp.exists(p) and shared.FUZZY_MATCH_IMAGE_NAME:
            if self._fuzzy_inpainted_list is None:
                if osp.exists(self.inpainted_dir()):
                    self._fuzzy_inpainted_list = find_all_imgs(self.inpainted_dir(), sort=True)
                else:
                    self._fuzzy_inpainted_list = []
            pidx = self.pagename2idx(imgname)
            if pidx < len(self._fuzzy_inpainted_list):
                return osp.join(self.inpainted_dir(), self._fuzzy_inpainted_list[pidx])
        return p
    
    def load_inpainted_by_imgname(self, imgname: str, scale_to_src: bool = True) -> np.ndarray:
        inpainted = None
        mp = self.get_inpainted_path(imgname, get_last_modified=True)
        if mp is not None and osp.exists(mp):
            inpainted = imread(mp)
            if imgname == self.current_img and self.img_array is not None:
                h, w = self.img_array.shape[:2]
            else:
                i = Image.open(osp.join(self.directory, imgname))
                h, w = i.height, i.width
            ih, iw = inpainted.shape[:2]
            if ih != h or iw != w:
                inpainted = Image.fromarray(inpainted).resize((w, h), resample=Image.Resampling.LANCZOS)
                inpainted = np.array(inpainted)
        return inpainted

    def get_result_path(self, imgname: str) -> str:
        ext = '.png'
        if pcfg is not None:
            if pcfg.imgsave_ext not in {'.jpg', '.png', '.webp', '.jxl'}:
                LOGGER.warning('invalid image saving ext in config.json')
            else:
                ext = pcfg.imgsave_ext
        return osp.join(self.result_dir(), osp.splitext(imgname)[0]+ext)
        
    def backup(self):
        raise NotImplementedError

    @property
    def is_empty(self):
        return len(self.pages) == 0

    @property
    def is_all_pages_no_text(self):
        return all([len(blklist) == 0 for blklist in self.pages.values()])

    @property
    def img_valid(self):
        return self.img_array is not None
    
    @property
    def mask_valid(self):
        return self.mask_array is not None

    @property
    def inpainted_valid(self):
        return self.inpainted_array is not None

    def set_next_img(self):
        if self.current_img is not None:
            next_idx = (self.current_idx + 1) % self.num_pages
            self.set_current_img(self.idx2pagename(next_idx))

    def set_prev_img(self):
        if self.current_img is not None:
            next_idx = (self.current_idx - 1 + self.num_pages) % self.num_pages
            self.set_current_img(self.idx2pagename(next_idx))

    def current_block_list(self) -> List[TextBlock]:
        if self.current_img is not None:
            assert self.current_img in self.pages
            return self.pages[self.current_img]
        else:
            return None

    def doc_path(self) -> str:
        return os.path.join(self.directory, self.proj_name() + ".docx")

    def doc_exist(self) -> bool:
        return osp.exists(self.doc_path())

    def dump_doc(self, delete_tmp_folder=True, fin_page_signal=None):
        
        cuts_dir = os.path.join(self.directory, "bubcuts")
        if os.path.exists(cuts_dir):
            shutil.rmtree(cuts_dir)
        os.mkdir(cuts_dir)
        
        document = Document()
        style = document.styles['Normal']
        font = style.font
        target_font = 'Arial'
        font.name = target_font
        for pagename, blklist in self.pages.items():
            imgpath = os.path.join(self.directory, pagename)
            
            cuts_path_list, cut_width_list = gen_ballon_cuts(cuts_dir, imgpath, blklist)
            paragraph = document.add_paragraph(pagename)
            paragraph.style = document.styles['Normal']
            table = document.add_table(rows=len(cuts_path_list), cols=2, style='Table Grid')

            for index, (cut_path, width) in enumerate(zip(cuts_path_list, cut_width_list)):
                run = table.cell(index, 0).paragraphs[0].add_run()
                run.style.font.name = target_font
                blk: TextBlock = blklist[index]
                bubdict = vars(blk).copy()
                bubdict["imgkey"] = pagename
                bubdict["rich_text"] = ''
                bubdict["text"] = blk.get_text()
                write_jpg_metadata(cut_path, metadata=json.dumps(bubdict, ensure_ascii=False, cls=TextBlkEncoder))
                run.add_picture(cut_path, width=Inches(width/96 * 0.85))
                table.cell(index, 1).text = bubdict["translation"]

            document.add_page_break()
            
            if fin_page_signal is not None:
                fin_page_signal.emit()
                # time.sleep(1)

        doc_path = self.doc_path()
        document.save(doc_path)
        if delete_tmp_folder:
            shutil.rmtree(cuts_dir)

    def dump_txt_path(self, dump_target, suffix):
        save_path = osp.join(self.directory, self.proj_name() + f'_{dump_target}{suffix}')
        return save_path

    def dump_txt(self, dump_target: str, suffix='.txt'):
        save_path = self.dump_txt_path(dump_target, suffix=suffix)
        text_all = []
        assert dump_target in {'source', 'translation'}
        assert suffix in {'.txt', '.md'}
        for page_name, blk_list in self.pages.items():
            text_in_page = ['### ' + page_name]
            for ii, blk in enumerate(blk_list):
                if dump_target == 'translation':
                    text = blk.translation.strip()
                elif dump_target == 'source':
                    text = blk.get_text().strip()
                text_in_page.append(f'{ii + 1}. {text}')
            text_all.append('\n\n'.join(text_in_page))
        with open(save_path, 'w', encoding='utf8') as f:
            f.write('\n\n\n'.join(text_all))

    def load_doc(self, doc_path, delete_tmp_folder=True, fin_page_signal=None):
        tmp_bubble_folder = osp.join(self.directory, 'img_folder')
        os.makedirs(tmp_bubble_folder, exist_ok=True)
        docx2txt.process(doc_path, tmp_bubble_folder)

        doc = docx.Document(doc_path)
        body_xml_str = doc._body._element.xml

        pages = {}
        effect_notices = set()
        bub_index = 0
        for tbl in re.findall(r'<w:tbl>(.*?)</w:tbl>', body_xml_str, re.DOTALL):
            for tr in re.findall(r'<w:tr(.*?)>(.*?)</w:tr>', tbl, re.DOTALL):
                if re.findall(r'<pic:cNvPr id=\"(.*?)\" name=\"(.*?)\"(.*?)>', tr[1]):
                    bub_index += 1
                    translation = ""
                    for paragraph in re.findall(r'<w:p(.*?)>(.*?)</w:p>', tr[1], re.DOTALL):
                        for wt in re.findall(r'<w:t>(.*?)</w:t>', paragraph[1], re.DOTALL):
                            translation += wt
                        translation += "\n"
                    translation = translation[:-1]
                    if len(translation) != 0 and translation[0] == "\n":
                        translation = translation[1:]


                    bubpath = os.path.join(tmp_bubble_folder, "image"+str(bub_index))
                    if osp.exists(bubpath+'.jpg'):
                        bubpath = bubpath + '.jpg'
                    else:
                        bubpath = bubpath + '.jpeg'

                    meta_dict = read_jpg_metadata(bubpath)
                    meta_dict["translation"] = translation
                    imgkey = meta_dict.pop("imgkey")
                    if not imgkey in pages:
                        pages[imgkey] = []
                    normalized, notices = normalize_textblock_effect_payload(
                        meta_dict
                    )
                    effect_notices.update(notices)
                    pages[imgkey].append(TextBlock(**normalized))
                    
                    if fin_page_signal is not None:
                        fin_page_signal.emit()

        warn_ignored_legacy_effects(effect_notices, 'document import')
        self.merge_from_proj_dict(pages)
        if delete_tmp_folder:
            shutil.rmtree(tmp_bubble_folder)

    def merge_from_proj_dict(self, tgt_dict: Dict) -> Dict:
        if self.pages is None:
            self.pages = {}
        src_dict = self.pages if self.pages is not None else {}
        key_lst = list(dict.fromkeys(list(src_dict.keys()) + list(tgt_dict.keys())))
        key_lst.sort()
        rst_dict = {}
        pagename2idx = {}
        idx2pagename = {}
        page_counter = 0
        for key in key_lst:
            if key in src_dict and not key in tgt_dict:
                rst_dict[key] = src_dict[key]
            else:
                rst_dict[key] = tgt_dict[key]
            pagename2idx[key] = page_counter
            idx2pagename[page_counter] = key
            page_counter += 1
        self.pages.clear()
        self.pages.update(rst_dict)
        self._pagename2idx = pagename2idx
        self._idx2pagename = idx2pagename        


def gen_ballon_cuts(cuts_dir: str, imgpath: str, blk_list: List[TextBlock], resize=True) -> Tuple[List[str], List[int]]:
    img = imread(imgpath)
    imgname = os.path.basename(imgpath)
    cuts_path_list = []
    cut_width_list = []
    for ii, blk in enumerate(blk_list):
        
        x, y, w, h = blk.bounding_rect()
        x, y = max(x, 0), max(y, 0)
        w = max(w, 1)
        h = max(h, 1)
        x1, y1, x2, y2 = int(x), int(y), int(x+w), int(y+h)

        cut_path = os.path.join(cuts_dir, f'{imgname}-{ii}.jpg')
        bub = img[y1:y2, x1:x2]
        max_width = 448

        if bub.shape[0] < 1 or bub.shape[1] < 1:
            emptyw = 60
            resized = np.full((emptyw, emptyw, 3), fill_value=0, dtype=np.uint8)
            width = emptyw
        else:
            # scale_percent = 60 # percent of original size
            scale_percent = min(1920 / img.shape[0], max_width / w)
            
            if scale_percent < 1:
                width = max(1, int(bub.shape[1] * scale_percent))
                height = max(1, int(bub.shape[0] * scale_percent))
                dim = (width, height)
                resized = cv2.resize(bub, dim, interpolation = cv2.INTER_AREA) if resize else bub
            else:
                width = w
                resized = bub

        imwrite(cut_path, resized, '.jpg')
        cuts_path_list.append(cut_path)
        cut_width_list.append(width)

    return cuts_path_list, cut_width_list
