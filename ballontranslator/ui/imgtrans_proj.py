import os, json, shutil, re, docx, docx2txt, piexif, cv2, time
from docx.shared import Inches
from docx import Document
import piexif.helper
import numpy as np
import os.path as osp
from typing import Tuple, Union, List, Dict

from utils.logger import logger as LOGGER
from utils.io_utils import find_all_imgs, imread, imwrite, NumpyEncoder
from dl.textdetector.textblock import TextBlock
from .misc import ImgnameNotInProjectException, ProjectLoadFailureException, ProjectDirNotExistException, ProjectNotSupportedException, TextBlkEncoder, KritaFile


def write_jpg_metadata(imgpath: str, metadata="a metadata"):
    exif_dict = {"Exif":{piexif.ExifIFD.UserComment: piexif.helper.UserComment.dump(metadata, encoding='unicode')}}
    exif_bytes = piexif.dump(exif_dict)
    piexif.insert(exif_bytes, imgpath)

def read_jpg_metadata(imgpath: str):
    exif_dict = piexif.load(imgpath)
    user_comment = piexif.helper.UserComment.load(exif_dict["Exif"][piexif.ExifIFD.UserComment])
    bubdict = json.loads(user_comment)
    return bubdict


class ProjImgTrans:

    def __init__(self, directory: str = None):
        self.type = 'imgtrans'
        self.directory: str = None
        self.pages: Dict[List[TextBlock]] = {}
        self._pagename2idx = {}
        self._idx2pagename = {}

        self.not_found_pages: Dict[List[TextBlock]] = {}
        self.new_pages: List[str] = []
        self.proj_path: str = None

        self.src_download_link: str = ''

        self.current_img: str = None
        self.img_array: np.ndarray = None
        self.mask_array: np.ndarray = None
        self.inpainted_array: np.ndarray = None
        if directory is not None:
            self.load(directory)

    def idx2pagename(self, idx: int) -> str:
        return self._idx2pagename[idx]

    def pagename2idx(self, pagename: str) -> int:
        if pagename in self.pages:
            return self._pagename2idx[pagename]
        return -1

    def proj_name(self) -> str:
        return self.type+'_'+osp.basename(self.directory)

    def load(self, directory: str, json_path: str = None) -> bool:
        self.directory = directory
        if json_path is None:
            self.proj_path = osp.join(self.directory, self.proj_name() + '.json')
        else:
            self.proj_path = json_path
        new_proj = False
        if not osp.exists(self.proj_path):
            new_proj = True
            self.new_project()
        else:
            try:
                with open(self.proj_path, 'r', encoding='utf8') as f:
                    proj_dict = json.loads(f.read())
            except Exception as e:
                raise ProjectLoadFailureException(e)
            self.load_from_dict(proj_dict)
        if not osp.exists(self.inpainted_dir()):
            os.makedirs(self.inpainted_dir())
        if not osp.exists(self.mask_dir()):
            os.makedirs(self.mask_dir())
        return new_proj

    def mask_dir(self):
        return osp.join(self.directory, 'mask')

    def inpainted_dir(self):
        return osp.join(self.directory, 'inpainted')

    def result_dir(self):
        return osp.join(self.directory, 'result')
    
    def init_properties(self, src_download_link: str = '', **kwargs):
        self.src_download_link = src_download_link

    def load_from_dict(self, proj_dict: dict):
        self.set_current_img(None)
        try:
            self.pages = {}
            self._pagename2idx = {}
            self._idx2pagename = {}
            self.not_found_pages = {}
            page_dict = proj_dict['pages']
            not_found_pages = list(page_dict.keys())
            found_pages = find_all_imgs(img_dir=self.directory, abs_path=False)
            for ii, imname in enumerate(found_pages):
                if imname in page_dict:
                    self.pages[imname] = [TextBlock(**blk_dict) for blk_dict in page_dict[imname]]
                    not_found_pages.remove(imname)
                else:
                    self.pages[imname] = []
                    self.new_pages.append(imname)
                self._pagename2idx[imname] = ii
                self._idx2pagename[ii] = imname
            for imname in not_found_pages:
                self.not_found_pages[imname] = [TextBlock(**blk_dict) for blk_dict in page_dict[imname]]
            self.init_properties(**proj_dict)
        except Exception as e:
            raise ProjectNotSupportedException(e)
        set_img_failed = False
        if 'current_img' in proj_dict:
            current_img = proj_dict['current_img']
            try:
                self.set_current_img(current_img)
            except ImgnameNotInProjectException:
                set_img_failed = True
        else:
            set_img_failed = True
            LOGGER.warning(f'{current_img} not found.')
        if set_img_failed:
            if len(self.pages) > 0:
                self.set_current_img_byidx(0)

    def load_from_json(self, json_path: str):
        old_dir = self.directory
        directory = osp.dirname(json_path)
        try:
            self.load(directory, json_path=json_path)
        except Exception as e:
            self.load(old_dir)
            raise ProjectLoadFailureException(e)

    def set_current_img(self, imgname: str):
        if imgname is not None:
            if imgname not in self.pages:
                raise ImgnameNotInProjectException
            self.current_img = imgname
            img_path = self.current_img_path()
            mask_path = self.mask_path()
            inpainted_path = self.inpainted_path()
            self.img_array = imread(img_path)
            im_h, im_w = self.img_array.shape[:2]
            if osp.exists(mask_path):
                self.mask_array = imread(mask_path, cv2.IMREAD_GRAYSCALE)
            else:
                self.mask_array = np.zeros((im_h, im_w), dtype=np.uint8)
            self.inpainted_array = imread(inpainted_path) if osp.exists(inpainted_path) else np.copy(self.img_array)
        else:
            self.current_img = None
            self.img_array = None
            self.mask_array = None
            self.inpainted_array = None

    def set_current_img_byidx(self, idx: int):
        num_pages = self.num_pages
        if idx < 0:
            idx = idx + self.num_pages
        if idx < 0 or idx > num_pages - 1:
            self.set_current_img(None)
        else:
            self.set_current_img(self.idx2pagename(idx))

    def get_blklist_byidx(self, idx: int) -> List[TextBlock]:
        return self.pages[self.idx2pagename(idx)]

    @property
    def num_pages(self) -> int:
        return len(self.pages)

    @property
    def current_idx(self) -> int:
        return self.pagename2idx(self.current_img)

    def new_project(self):
        if not osp.exists(self.directory):
            raise ProjectDirNotExistException
        self.set_current_img(None)
        imglist = find_all_imgs(self.directory, abs_path=False)
        self.pages = {}
        self._pagename2idx = {}
        self._idx2pagename = {}
        for ii, imgname in enumerate(imglist):
            self.pages[imgname] = []
            self._pagename2idx[imgname] = ii
            self._idx2pagename[ii] = imgname
        self.init_properties()
        self.set_current_img_byidx(0)
        self.save()
        
    def save(self):
        if not osp.exists(self.directory):
            raise ProjectDirNotExistException
        with open(self.proj_path, "w", encoding="utf-8") as f:
            f.write(json.dumps(self.to_dict(), ensure_ascii=False, cls=TextBlkEncoder))

    def to_dict(self) -> Dict:
        pages = self.pages.copy()
        pages.update(self.not_found_pages)
        return {
            'directory': self.directory,
            'pages': pages,
            'current_img': self.current_img,
            'src_download_link': self.src_download_link
        }

    def read_img(self, imgname: str) -> np.ndarray:
        if imgname not in self.pages:
            raise ImgnameNotInProjectException
        return imread(osp.join(self.directory, imgname))

    def save_mask(self, img_name, mask: np.ndarray):
        imwrite(self.get_mask_path(img_name), mask)

    def save_inpainted(self, img_name, inpainted: np.ndarray):
        imwrite(self.get_inpainted_path(img_name), inpainted)

    def current_img_path(self) -> str:
        if self.current_img is None:
            return None
        return osp.join(self.directory, self.current_img)

    def mask_path(self) -> str:
        if self.current_img is None:
            return None
        return self.get_mask_path(self.current_img)

    def inpainted_path(self) -> str:
        if self.current_img is None:
            return None
        return self.get_inpainted_path(self.current_img)

    def get_mask_path(self, imgname: str = None) -> str:
        if imgname is None:
            imgname = self.current_img
        return osp.join(self.mask_dir(), osp.splitext(imgname)[0]+'.png')

    def get_inpainted_path(self, imgname: str = None) -> str:
        if imgname is None:
            imgname = self.current_img
        return osp.join(self.inpainted_dir(), osp.splitext(imgname)[0]+'.png')

    def get_result_path(self, imgname: str) -> str:
        return osp.join(self.result_dir(), osp.splitext(imgname)[0]+'.png')
        
    def backup(self):
        raise NotImplementedError

    @property
    def is_empty(self):
        return len(self.pages) == 0

    @property
    def img_valid(self):
        return self.img_array is not None
    
    @property
    def mask_valid(self):
        return self.mask_array is not None

    @property
    def inpainted_valid(self):
        return self.inpainted_array is not None

    def set_next_img(self):
        if self.current_img is not None:
            next_idx = (self.current_idx + 1) % self.num_pages
            self.set_current_img(self.idx2pagename(next_idx))

    def set_prev_img(self):
        if self.current_img is not None:
            next_idx = (self.current_idx - 1 + self.num_pages) % self.num_pages
            self.set_current_img(self.idx2pagename(next_idx))

    def current_block_list(self) -> List[TextBlock]:
        if self.current_img is not None:
            assert self.current_img in self.pages
            return self.pages[self.current_img]
        else:
            return None

    def doc_path(self) -> str:
        return os.path.join(self.directory, self.proj_name() + ".docx")

    def doc_exist(self) -> bool:
        return osp.exists(self.doc_path())

    def dump_doc(self, delete_tmp_folder=True, fin_page_signal=None):
        
        cuts_dir = os.path.join(self.directory, "bubcuts")
        if os.path.exists(cuts_dir):
            shutil.rmtree(cuts_dir)
        os.mkdir(cuts_dir)
        
        document = Document()
        style = document.styles['Normal']
        font = style.font
        target_font = 'Arial'
        font.name = target_font
        for pagename, blklist in self.pages.items():
            imgpath = os.path.join(self.directory, pagename)
            
            cuts_path_list, cut_width_list = gen_ballon_cuts(cuts_dir, imgpath, blklist)
            paragraph = document.add_paragraph(pagename)
            paragraph.style = document.styles['Normal']
            table = document.add_table(rows=len(cuts_path_list), cols=2, style='Table Grid')

            for index, (cut_path, width) in enumerate(zip(cuts_path_list, cut_width_list)):
                run = table.cell(index, 0).paragraphs[0].add_run()
                run.style.font.name = target_font
                blk: TextBlock = blklist[index]
                bubdict = vars(blk).copy()
                bubdict["imgkey"] = pagename
                bubdict["rich_text"] = ''
                bubdict["text"] = blk.get_text()
                write_jpg_metadata(cut_path, metadata=json.dumps(bubdict, ensure_ascii=False, cls=TextBlkEncoder))
                run.add_picture(cut_path, width=Inches(width/96 * 0.85))
                table.cell(index, 1).text = bubdict["translation"]

            document.add_page_break()
            
            if fin_page_signal is not None:
                fin_page_signal.emit()
                # time.sleep(1)

        doc_path = self.doc_path()
        document.save(doc_path)
        if delete_tmp_folder:
            shutil.rmtree(cuts_dir)


    def load_doc(self, doc_path, delete_tmp_folder=True, fin_page_signal=None):

        tmp_bubble_folder = osp.join(self.directory, 'img_folder')
        os.makedirs(tmp_bubble_folder, exist_ok=True)
        docx2txt.process(doc_path, tmp_bubble_folder)

        doc = docx.Document(doc_path)
        body_xml_str = doc._body._element.xml

        pages = {}
        bub_index = 0
        for tbl in re.findall(r'<w:tbl>(.*?)</w:tbl>', body_xml_str, re.DOTALL):
            for tr in re.findall(r'<w:tr(.*?)>(.*?)</w:tr>', tbl, re.DOTALL):
                if re.findall(r'<pic:cNvPr id=\"0\" name=\"(.*?)\"/>', tr[1]):
                    bub_index += 1
                    translation = ""
                    for paragraph in re.findall(r'<w:p(.*?)>(.*?)</w:p>', tr[1], re.DOTALL):
                        for wt in re.findall(r'<w:t>(.*?)</w:t>', paragraph[1], re.DOTALL):
                            translation += wt
                        translation += "\n"
                    translation = translation[:-1]
                    if len(translation) != 0 and translation[0] == "\n":
                        translation = translation[1:]
                    bubpath = os.path.join(tmp_bubble_folder, "image"+str(bub_index)+".jpg")

                    meta_dict = read_jpg_metadata(bubpath)
                    meta_dict["translation"] = translation
                    imgkey = meta_dict.pop("imgkey")
                    if not imgkey in pages:
                        pages[imgkey] = []
                    pages[imgkey].append(TextBlock(**meta_dict))
                    
                    if fin_page_signal is not None:
                        fin_page_signal.emit()

        self.merge_from_proj_dict(pages)
        if delete_tmp_folder:
            shutil.rmtree(tmp_bubble_folder)

    def merge_from_proj_dict(self, tgt_dict: Dict) -> Dict:
        if self.pages is None:
            self.pages = {}
        src_dict = self.pages if self.pages is not None else {}
        key_lst = list(dict.fromkeys(list(src_dict.keys()) + list(tgt_dict.keys())))
        key_lst.sort()
        rst_dict = {}
        pagename2idx = {}
        idx2pagename = {}
        page_counter = 0
        for key in key_lst:
            if key in src_dict and not key in tgt_dict:
                rst_dict[key] = src_dict[key]
            else:
                rst_dict[key] = tgt_dict[key]
            pagename2idx[key] = page_counter
            idx2pagename[page_counter] = key
            page_counter += 1
        self.pages.clear()
        self.pages.update(rst_dict)
        self._pagename2idx = pagename2idx
        self._idx2pagename = idx2pagename        

    def dump_kra(self, text=1, original=1, mask=1, inpainted=1, result=1, kra_dir=None, copy_src_to=None,fin_page_signal=None):
        # TODO
        # if text is 1 put textboxes in one layer, if 2 separate them into diffent layers.
        # if mask is 2 or something, crop/merge inpainted_img with mask so it becomes one layer.
        # saving only one image at a time (current image).
        # copy_src_to is (planned) for copying the mask, inpainted, and result dirs
        # to another dir so it's referenced from there and not from the imgtrans project's files
        
        if kra_dir == None:
            kra_dir = os.path.join(self.directory, 'krita')    
        
        os.makedirs(kra_dir, exist_ok=True)
        
        self.kra_path = kra_dir

        def layer_src_path(path, kra_dir=kra_dir):
            return os.path.relpath(path, kra_dir)

        for pagename, blklist in self.pages.items():
            kra_doc = KritaFile()
            img = self.read_img(pagename)

            if text == 1:
                # TODO
                pass
            if result == 1:
                result_path = self.get_result_path(pagename)
                src_path = layer_src_path(result_path)
                kra_doc.add_file_layer('Result Layer', src_path)
            if inpainted == 1:
                inpaint_path = self.get_inpainted_path(pagename)
                src_path = layer_src_path(inpaint_path)
                kra_doc.add_file_layer('Inpaint Layer', src_path)
            if mask == 1:
                mask_path = self.get_mask_path(pagename)
                src_path = layer_src_path(mask_path)
                kra_doc.add_file_layer('Mask Layer', src_path)
            if original == 1:
                img_path = os.path.join(self.directory, pagename)
                src_path = layer_src_path(img_path)
                kra_doc.add_file_layer('Original Image', src_path)
            
            kra_file = pagename.split('.')[0] + '.kra'
            kra_file_path = os.path.join(kra_dir, kra_file)
            kra_doc.write_kra_file(thumbnail=img, path=kra_file_path)

            if fin_page_signal is not None:
                fin_page_signal.emit()

def gen_ballon_cuts(cuts_dir: str, imgpath: str, blk_list: List[TextBlock], resize=True) -> Tuple[List[str], List[int]]:
    img = imread(imgpath)
    imgname = os.path.basename(imgpath)
    cuts_path_list = []
    cut_width_list = []
    for ii, blk in enumerate(blk_list):
        
        x, y, w, h = blk.bounding_rect()
        x, y = max(x, 0), max(y, 0)
        w = max(w, 1)
        h = max(h, 1)
        x1, y1, x2, y2 = int(x), int(y), int(x+w), int(y+h)

        cut_path = os.path.join(cuts_dir, f'{imgname}-{ii}.jpg')
        bub = img[y1:y2, x1:x2]
        # scale_percent = 60 # percent of original size
        scale_percent = 1280 / img.shape[0]
        width = max(1, int(bub.shape[1] * scale_percent))
        height = max(1, int(bub.shape[0] * scale_percent))
        dim = (width, height)
        resized = cv2.resize(bub, dim, interpolation = cv2.INTER_AREA) if resize else bub

        imwrite(cut_path, resized, '.jpg')
        cuts_path_list.append(cut_path)
        cut_width_list.append(width)

    return cuts_path_list, cut_width_list



